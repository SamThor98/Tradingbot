"""Generate a plain-language Word overview of the TradingBot system."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "TradingBot_System_Overview.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_document() -> Document:
    doc = Document()

    # Title page
    title = doc.add_heading("TradingBot", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("A Plain-English Guide to How the System Works")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()
    date_p = doc.add_paragraph("June 2026")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Introduction ---
    add_heading(doc, "What Is This System?", 1)
    add_para(
        doc,
        "TradingBot is an automated stock-trading assistant that connects to your Charles Schwab "
        "brokerage account. Think of it as a tireless research analyst that watches thousands of "
        "stocks every day, finds ones that look ready to rise, checks them from multiple angles, "
        "and—only if they pass every safety check—suggests or places a trade on your behalf.",
    )
    add_para(
        doc,
        "It is built around a well-known style of investing called momentum trading: buying stocks "
        "that are already going up, in strong sectors, before they break out to new highs. The system "
        "does not guess at random. It follows a structured recipe that professional traders have used "
        "for decades, then adds modern checks (earnings data, financial health scores, and machine-learning "
        "estimates) on top.",
    )
    add_para(
        doc,
        "Important: This software is a tool, not financial advice. You remain responsible for every "
        "trade. The system includes multiple safety layers and can run in “paper trading” mode where "
        "no real money moves.",
        bold=True,
    )

    # --- Big Picture ---
    add_heading(doc, "The Big Picture", 1)
    add_para(
        doc,
        "At the highest level, the system does four things in a loop:",
    )
    add_bullets(
        doc,
        [
            "Collect price and market data (mostly from Schwab, with backups from other sources).",
            "Scan for promising stock setups using technical and fundamental filters.",
            "Apply safety rules so no single bad trade can blow up the account.",
            "Execute approved trades and learn from the results over time.",
        ],
    )
    add_para(doc, "Here is the daily flow in simple terms:")
    add_para(
        doc,
        "Morning → Check that Schwab connections are healthy and the overall market looks favorable.\n"
        "During the day → Run scans, rank candidates, send alerts.\n"
        "When a signal looks strong → Queue it for your approval (or auto-trade if you enable that).\n"
        "After a trade → Attach a protective stop-loss and log everything.\n"
        "End of day → Review what happened and update internal “learning” files.",
    )

    # --- Schwab Connection ---
    add_heading(doc, "Connecting to Schwab", 1)
    add_para(
        doc,
        "The system talks to Schwab through two separate, secure logins—on purpose:",
    )
    add_bullets(
        doc,
        [
            "Market session — reads stock prices, charts, and quotes. This is read-only market data.",
            "Account session — sees your balances and places orders. This is the “hands on the steering wheel” connection.",
        ],
    )
    add_para(
        doc,
        "Splitting these two roles is a safety feature. A bug in the scanning code cannot accidentally "
        "place an order, because order placement uses a completely different credential. You authorize "
        "both sessions once through your browser (standard OAuth, the same kind of “Log in with Google” "
        "flow many apps use). Tokens are stored encrypted on your machine.",
    )

    # --- What stocks it looks for ---
    add_heading(doc, "What Kind of Stocks Does It Look For?", 1)

    add_heading(doc, "Stage 2 — Stocks in a Strong Uptrend", 2)
    add_para(
        doc,
        "Market analyst Stan Weinstein described four “stages” every stock goes through: basing, rising, "
        "topping, and falling. TradingBot focuses on Stage 2—the rising phase—when institutions are "
        "quietly accumulating shares before the public notices.",
    )
    add_para(doc, "A stock must pass two simple checks:")
    add_bullets(
        doc,
        [
            "Its price is near its 52-week high (within about 15% by default)—it is a leader, not a laggard.",
            "Its 200-day moving average has been rising for at least 20 days—the long-term trend is up.",
        ],
    )

    add_heading(doc, "VCP — Volume Contraction Pattern", 2)
    add_para(
        doc,
        "After a stock is in an uptrend, traders look for a “coiled spring” pattern: the price tightens "
        "into a smaller and smaller range while trading volume dries up. That usually means sellers are "
        "exhausted. When volume picks back up and price breaks above the range, it often signals the "
        "next leg higher. The system looks for several consecutive days of below-average volume before "
        "flagging a candidate.",
    )

    add_heading(doc, "Sector Strength", 2)
    add_para(
        doc,
        "Even a great-looking stock in a weak industry tends to underperform. Optionally, the scanner "
        "only keeps stocks whose sector is beating the broad market (SPY). Buying strength within "
        "strength improves the odds.",
    )

    add_heading(doc, "Market Regime Gate", 2)
    add_para(
        doc,
        "Before scanning individual stocks, the system checks the overall market. By default, if the "
        "S&P 500 (SPY) is below its 200-day average—a classic sign of a bear market—it will not run "
        "new buy scans at all. This “fail closed” approach avoids fighting the tide.",
    )

    # --- Two-stage scanner ---
    add_heading(doc, "How Scanning Works (Two Stages)", 1)
    add_para(
        doc,
        "Scanning thousands of stocks with heavy analysis on every one would be slow and expensive. "
        "So the system uses a funnel—two stages, like a job interview with a phone screen first and "
        "a full panel later.",
    )

    add_heading(doc, "Stage A — The Quick Filter", 2)
    add_para(doc, "Fast, cheap checks run in parallel on a large universe of tickers:")
    add_bullets(
        doc,
        [
            "Is it in Stage 2 (strong uptrend)?",
            "Does it show a Volume Contraction Pattern?",
            "Is its sector outperforming (if that filter is on)?",
        ],
    )
    add_para(
        doc,
        "Only the survivors move to Stage B. Typically a few dozen names make the shortlist.",
    )

    add_heading(doc, "Stage B — Deep Research", 2)
    add_para(doc, "Each shortlisted stock gets a thorough look:")
    add_bullets(
        doc,
        [
            "Post-earnings drift (PEAD) — did recent earnings surprise positively? Stocks often drift in the direction of the surprise.",
            "Forensic accounting — automated checks for accounting red flags (Sloan ratio, Beneish M-score, Altman Z-score). Think of this as a smoke detector for financial statements.",
            "SEC filings — reads public filing hints for risk language and score changes.",
            "Advisory model — a trained statistical model estimates the probability the stock will be higher in ~10 days, with High / Medium / Low confidence bands.",
            "MiroFish simulation — several virtual “personas” (momentum trader, value investor, mean-reversion specialist, news watcher, etc.) each vote on the stock. Their votes are weighted by how well each persona has performed recently in the current market regime.",
            "Prediction markets (optional) — can overlay odds from platforms like Polymarket when liquidity is good enough.",
            "Quality gates — weak or contradictory signals get filtered out.",
            "Final ranking — everything is scored and sorted; only the top candidates are surfaced.",
        ],
    )
    add_para(
        doc,
        "Every scan also returns a diagnostics report: how many stocks failed at each step and why. "
        "If you ever wonder “why no signals today?”, the dashboard shows exactly which filter blocked them.",
    )

    # --- Safety ---
    add_heading(doc, "Safety Rules (Guardrails)", 1)
    add_para(
        doc,
        "Automated trading without limits is dangerous. TradingBot wraps every order in hard rules:",
    )
    add_bullets(
        doc,
        [
            "Account size cap — will not let total exposure exceed a configured maximum (default $500,000).",
            "Per-stock cap — limits how much can go into any single ticker (default $50,000).",
            "Daily trade limit — caps how many trades can fire in one day (default 20).",
            "Sector cap — limits how much of the portfolio can sit in one industry.",
            "Trailing stop-loss — when you buy, a protective stop (default ~7% below entry) is attached automatically so a bad trade has a defined maximum loss.",
            "Adaptive stops — stop distance can widen or tighten based on how volatile the stock is (ATR-based sizing).",
            "Kill switch — one environment variable instantly halts all new buys platform-wide.",
            "Data quality checks — if price quotes are stale or conflicting, new risk-increasing trades can be blocked while exits are still allowed.",
            "Circuit breaker — if Schwab’s API fails repeatedly, the system pauses and waits before retrying.",
        ],
    )
    add_para(
        doc,
        "You can also run in shadow or paper mode: the system computes what it would do and logs the "
        "decision, but never sends a real order.",
    )

    # --- Execution ---
    add_heading(doc, "Placing Trades", 1)
    add_para(doc, "When a signal clears every filter and guardrail:")
    add_bullets(
        doc,
        [
            "Position size is calculated based on your risk settings and the stock’s volatility.",
            "An order is submitted through the Schwab account session.",
            "On fill confirmation, a stop-loss order is attached.",
            "The trade is logged to the database and an alert is sent (Discord, dashboard).",
        ],
    )
    add_para(
        doc,
        "By default on the local dashboard, trades land in a pending queue—you approve or reject each "
        "one before it goes live. This human-in-the-loop step is recommended until you trust the system’s "
        "behavior.",
    )
    add_para(doc, "Optional “plugins” add extra behavior (each rolls out slowly: off → shadow → live):")
    add_bullets(
        doc,
        [
            "Execution quality — checks spread and slippage before sending.",
            "Exit manager — partial take-profits, move stop to breakeven, time-based exits.",
            "Event risk — blocks new entries around earnings or macro events.",
            "Regime sizing — scales position size up or down based on market conditions.",
            "Correlation guard — avoids loading up on stocks that move together.",
        ],
    )

    # --- Learning ---
    add_heading(doc, "How the System Learns", 1)
    add_para(
        doc,
        "TradingBot is not static. It records outcomes and tries to improve calibration over time—"
        "separate from whether you made money on any given day.",
    )

    add_heading(doc, "Self-Study", 2)
    add_para(
        doc,
        "After enough completed round-trip trades (buy then sell), the system analyzes results by "
        "conviction band and sector. It may suggest raising the minimum conviction threshold for "
        "future alerts. Runs automatically at 4:00 PM Eastern.",
    )

    add_heading(doc, "Hypothesis Ledger", 2)
    add_para(
        doc,
        "Every signal that triggers an alert can be recorded as a “hypothesis”: ticker, predicted "
        "direction, reference price. Later, a scoring script checks what actually happened at 1-day, "
        "5-day, and 20-day horizons. This measures decision quality—not just profit and loss.",
    )

    add_heading(doc, "Feature Store & Evolution", 2)
    add_para(
        doc,
        "Scan features and outcomes are stored over time. Analysis scripts can propose small tuning "
        "changes (threshold adjustments) and test them as challengers against the current champion "
        "settings before anything goes live.",
    )

    # --- Dashboard ---
    add_heading(doc, "The Web Dashboard", 1)
    add_para(
        doc,
        "A browser-based control panel runs on your computer (default: http://127.0.0.1:8000). "
        "It is organized around your daily workflow:",
    )
    add_bullets(
        doc,
        [
            "Today — system health, scan blockers, run scans, approve pending trades.",
            "Research — quick ticker check, backtests, SEC comparison, portfolio view.",
            "System — deep health checks, validation status, calibration metrics.",
            "Settings — Schwab connection, live-trading controls, risk presets.",
        ],
    )
    add_para(
        doc,
        "Display modes (Simple / Standard / Pro) hide or show advanced columns and panels so beginners "
        "are not overwhelmed. Sensitive actions like approving trades require your API key.",
    )

    # --- Discord ---
    add_heading(doc, "Discord Alerts", 1)
    add_para(
        doc,
        "If you connect a Discord webhook, the bot sends color-coded notifications:",
    )
    add_bullets(
        doc,
        [
            "Heartbeat — “I’m alive” morning check.",
            "Signal — new candidate with conviction and score.",
            "Order filled — trade executed successfully.",
            "Guardrail block — tried to trade but a safety rule stopped it.",
            "Crash / error — something broke and needs attention.",
        ],
    )
    add_para(
        doc,
        "High-conviction, high-score signals can @mention you so you do not miss them. You can also "
        "trigger a scan on demand with a /scan slash command in Discord.",
    )

    # --- Deployment modes ---
    add_heading(doc, "Two Ways to Run It", 1)

    add_heading(doc, "Local (Single User)", 2)
    add_bullets(
        doc,
        [
            "Runs on your PC or a private server.",
            "Uses a simple SQLite database.",
            "Protected by an API key you set.",
            "Scheduled tasks (morning heartbeat, afternoon self-study) run via a lightweight scheduler.",
            "Best for personal use.",
        ],
    )

    add_heading(doc, "SaaS (Multi-Tenant)", 2)
    add_bullets(
        doc,
        [
            "Cloud-hosted version for multiple paying users.",
            "Uses PostgreSQL, Redis, and background workers (Celery).",
            "Users log in with Supabase authentication; billing via Stripe.",
            "Each tenant’s data is isolated.",
            "Same scanning and guardrail logic, scaled for production.",
        ],
    )

    # --- Backtesting ---
    add_heading(doc, "Backtesting", 1)
    add_para(
        doc,
        "Before trusting live signals, you can replay history. The backtest engine uses the same Stage A "
        "and Stage B rules as live scanning, applied to past market data across different market eras "
        "(bull markets, bear markets, volatile chop, crash recoveries). Results include profit factor, "
        "win rate, and per-era breakdowns so you can see whether the strategy holds up in bad times—not "
        "just good ones.",
    )

    # --- Summary diagram as text ---
    add_heading(doc, "End-to-End Flow (Summary)", 1)
    add_para(
        doc,
        "Market Data (Schwab + backups)\n"
        "    ↓\n"
        "Regime check (is the broad market healthy?)\n"
        "    ↓\n"
        "Stage A scan (Stage 2 + VCP + sector)\n"
        "    ↓\n"
        "Stage B enrichment (earnings, accounting, SEC, advisory model, MiroFish, quality gates)\n"
        "    ↓\n"
        "Rank & alert (Discord + dashboard queue)\n"
        "    ↓\n"
        "Your approval (optional but recommended)\n"
        "    ↓\n"
        "Guardrails → Order → Stop-loss → Log\n"
        "    ↓\n"
        "Self-study & hypothesis scoring (learn for next time)",
    )

    # --- Disclaimer ---
    add_heading(doc, "Important Disclaimers", 1)
    add_bullets(
        doc,
        [
            "This software does not provide investment advice.",
            "Past backtest or paper-trading performance does not guarantee future results.",
            "All trading involves risk of loss, including loss of principal.",
            "Automated systems can fail due to bugs, data errors, or broker outages.",
            "You should review the legal disclosures in the dashboard footer and consult qualified professionals before trading.",
            "Schwab and other third-party names are trademarks of their respective owners; this project is not affiliated with or endorsed by them.",
        ],
    )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(
        "Document generated from the TradingBot codebase and wiki. "
        "For technical details, see schwab_skill/README.md and wiki/index.md."
    )
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
